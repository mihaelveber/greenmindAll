"""
Document Parser - Extract text from various file formats for AI processing
Supported formats: PDF, Word, Excel, CSV, TXT, Images (OCR)
"""

import logging
from typing import Tuple

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_FORMATS = {
    'pdf': 'application/pdf',
    'doc': 'application/msword',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xls': 'application/vnd.ms-excel',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'csv': 'text/csv',
    'txt': 'text/plain',
    'jpg': 'image/jpeg',
    'jpeg': 'image/jpeg',
    'png': 'image/png',
}


def get_supported_formats_message() -> str:
    """Return user-friendly message about supported formats"""
    formats = list(SUPPORTED_FORMATS.keys())
    return f"Supported formats: {', '.join(formats).upper()}"


def extract_text_from_pdf_with_ocr(file_path: str) -> str:
    """
    Extract text from image-based PDF using OCR
    Converts PDF pages to images and applies OCR to each page
    """
    try:
        from pdf2image import convert_from_path
        from PIL import Image
        import pytesseract

        logger.info(f'Converting PDF to images for OCR: {file_path}')

        # Convert PDF to list of images (one per page)
        images = convert_from_path(file_path)

        text_content = []
        for page_num, image in enumerate(images, start=1):
            logger.info(f'Running OCR on page {page_num}/{len(images)}')

            # Extract text using OCR
            text = pytesseract.image_to_string(image)

            if text.strip():
                text_content.append(f"--- Page {page_num} ---\n{text}")

        if not text_content:
            return "[No text detected in PDF using OCR]"

        logger.info(f'Successfully extracted text from {len(images)} pages using OCR')
        return '\n\n'.join(text_content)

    except ImportError as e:
        logger.error(f'OCR dependencies not available: {str(e)}')
        raise Exception('OCR not available for image-based PDFs. Required: pdf2image and pytesseract')

    except Exception as e:
        logger.error(f'PDF OCR extraction failed: {str(e)}')
        if 'tesseract' in str(e).lower():
            raise Exception('Tesseract OCR engine not installed. Required for image-based PDFs.')
        raise Exception(f'Failed to extract text from PDF using OCR: {str(e)}')


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file (with OCR fallback for image-based PDFs)"""
    try:
        import PyPDF2

        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            logger.info(f'Extracting text from PDF: {num_pages} pages')

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")

        # If no text was extracted, try OCR (image-based PDF)
        if not text_content:
            logger.info('No text extracted with PyPDF2, attempting OCR on PDF pages')
            return extract_text_from_pdf_with_ocr(file_path)

        return '\n\n'.join(text_content)

    except Exception as e:
        logger.error(f'PDF extraction failed: {str(e)}')
        raise Exception(f'Failed to extract text from PDF: {str(e)}')


def extract_text_from_word(file_path: str) -> str:
    """Extract text from Word document (.docx)"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_content = []
        
        logger.info(f'Extracting text from Word document: {len(doc.paragraphs)} paragraphs')
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)
        
        return '\n\n'.join(text_content)
    
    except Exception as e:
        logger.error(f'Word extraction failed: {str(e)}')
        raise Exception(f'Failed to extract text from Word document: {str(e)}')


def extract_text_from_excel(file_path: str) -> str:
    """Extract text from Excel file (.xlsx, .xls)"""
    try:
        import pandas as pd
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        text_content = []
        
        logger.info(f'Extracting text from Excel: {len(excel_file.sheet_names)} sheets')
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Convert DataFrame to text with better table preservation
            sheet_text = f"\n\n{'='*80}\n=== SHEET: {sheet_name} ===\n{'='*80}\n\n"
            
            # Add column headers with clear separator
            headers = ' | '.join(str(col) for col in df.columns)
            sheet_text += headers + '\n' + '-' * min(len(headers), 120) + '\n'
            
            # Add rows (limit to avoid too much data) with clear row separation
            max_rows = 1000
            for idx, row in df.head(max_rows).iterrows():
                row_text = ' | '.join(str(val) if pd.notna(val) else '' for val in row.values)
                sheet_text += row_text + '\n'
            
            if len(df) > max_rows:
                sheet_text += f'\n[... {len(df) - max_rows} more rows truncated ...]\n'
            
            # Add end marker to preserve sheet boundaries
            sheet_text += f"\n{'='*80}\n=== END OF SHEET: {sheet_name} ===\n{'='*80}\n\n"
            
            text_content.append(sheet_text)
        
        return '\n\n'.join(text_content)
    
    except Exception as e:
        logger.error(f'Excel extraction failed: {str(e)}')
        raise Exception(f'Failed to extract text from Excel: {str(e)}')


def extract_text_from_csv(file_path: str) -> str:
    """Extract text from CSV file"""
    try:
        import pandas as pd
        
        df = pd.read_csv(file_path)
        
        logger.info(f'Extracting text from CSV: {len(df)} rows, {len(df.columns)} columns')
        
        # Convert to text representation
        text_content = "=== CSV Data ===\n\n"
        
        # Add column headers
        headers = ' | '.join(str(col) for col in df.columns)
        text_content += headers + '\n' + '-' * len(headers) + '\n'
        
        # Add rows (limit to avoid too much data)
        max_rows = 1000
        for idx, row in df.head(max_rows).iterrows():
            row_text = ' | '.join(str(val) if pd.notna(val) else '' for val in row.values)
            text_content += row_text + '\n'
        
        if len(df) > max_rows:
            text_content += f'\n[... {len(df) - max_rows} more rows truncated ...]'
        
        return text_content
    
    except Exception as e:
        logger.error(f'CSV extraction failed: {str(e)}')
        raise Exception(f'Failed to extract text from CSV: {str(e)}')


def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using OCR (pytesseract)"""
    try:
        from PIL import Image
        import pytesseract
        
        logger.info('Extracting text from image using OCR')
        
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        
        if not text.strip():
            return "[No text detected in image]"
        
        return text
    
    except Exception as e:
        logger.error(f'Image OCR extraction failed: {str(e)}')
        # OCR can fail if tesseract not installed - provide helpful message
        if 'tesseract' in str(e).lower():
            raise Exception('OCR not available. Text extraction from images requires Tesseract installation.')
        raise Exception(f'Failed to extract text from image: {str(e)}')


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f'Text file extraction failed: {str(e)}')
            raise Exception(f'Failed to read text file: {str(e)}')


def parse_document(file_path: str, file_name: str) -> Tuple[str, str]:
    """
    Parse document and extract text content
    
    Returns:
        Tuple[str, str]: (extracted_text, format_info)
    
    Raises:
        ValueError: If file format is not supported
        Exception: If extraction fails
    """
    # Get file extension using os.path.splitext to handle complex filenames
    import os
    _, file_ext = os.path.splitext(file_name.lower())
    file_ext = file_ext.lstrip('.')  # Remove leading dot
    
    # Check if format is supported
    if file_ext not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Unsupported file format: .{file_ext}. {get_supported_formats_message()}"
        )
    
    logger.info(f'Parsing document: {file_name} (format: {file_ext})')
    
    # Extract text based on file type
    try:
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_path)
            format_info = 'PDF'
        
        elif file_ext in ['doc', 'docx']:
            text = extract_text_from_word(file_path)
            format_info = 'Word Document'
        
        elif file_ext in ['xls', 'xlsx']:
            text = extract_text_from_excel(file_path)
            format_info = 'Excel Spreadsheet'
        
        elif file_ext == 'csv':
            text = extract_text_from_csv(file_path)
            format_info = 'CSV Data'
        
        elif file_ext in ['jpg', 'jpeg', 'png']:
            text = extract_text_from_image(file_path)
            format_info = 'Image (OCR)'
        
        elif file_ext == 'txt':
            text = extract_text_from_txt(file_path)
            format_info = 'Plain Text'
        
        else:
            raise ValueError(f"Parser not implemented for .{file_ext}")
        
        # Validate extracted text
        if not text or not text.strip():
            raise Exception(f"No text content extracted from {file_name}")
        
        logger.info(f'Successfully extracted {len(text)} characters from {file_name}')
        
        return text, format_info
    
    except ValueError:
        # Re-raise ValueError (unsupported format)
        raise
    
    except Exception as e:
        logger.error(f'Document parsing failed for {file_name}: {str(e)}')
        raise Exception(f'Failed to parse {file_name}: {str(e)}')


def is_supported_format(file_name: str) -> bool:
    """Check if file format is supported"""
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    return file_ext in SUPPORTED_FORMATS
