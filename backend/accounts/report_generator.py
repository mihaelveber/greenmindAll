"""
ESRS Report Generator - PDF and Word Export
Generates professional branded reports with AI-analyzed styling
"""

import logging
from typing import Optional, List
from io import BytesIO
import base64
from datetime import datetime
from pathlib import Path
import os

# HTML/CSS to PDF
from weasyprint import HTML, CSS
from jinja2 import Template, Environment, FileSystemLoader

# Legacy Word Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

# Brand Analyzer
from .brand_analyzer import BrandAnalyzer

logger = logging.getLogger(__name__)


class ESRSReportGenerator:
    """Generate branded PDF and Word reports for ESRS disclosures using AI-analyzed styling"""

    def __init__(self, user, standard_id: Optional[int] = None, disclosure_ids: Optional[List[int]] = None):
        self.user = user
        self.standard_id = standard_id
        self.disclosure_ids = disclosure_ids
        self.brand_analyzer = BrandAnalyzer()

        # Setup Jinja2 template environment
        template_dir = Path(__file__).parent / 'templates' / 'reports'
        self.jinja_env = Environment(loader=FileSystemLoader(str(template_dir)))

    def _get_brand_style(self) -> tuple:
        """
        Get or generate brand style for the user
        Returns (style_guide dict, css string, logo_path)
        """
        # Check if user has brand style already
        if self.user.brand_style and self.user.company_logo:
            style_guide = self.user.brand_style
            css = self.brand_analyzer.generate_css_from_style(style_guide)
            logo_path = self.user.company_logo.path if hasattr(self.user.company_logo, 'path') else None
            return style_guide, css, logo_path

        # Check if user has logo but no style - analyze it
        if self.user.company_logo:
            try:
                logo_path = self.user.company_logo.path if hasattr(self.user.company_logo, 'path') else None
                if logo_path and os.path.exists(logo_path):
                    logger.info(f"Analyzing logo for user {self.user.email}")
                    style_guide = self.brand_analyzer.analyze_logo(logo_path)

                    # Save style guide to user
                    self.user.brand_style = style_guide
                    self.user.save(update_fields=['brand_style'])

                    css = self.brand_analyzer.generate_css_from_style(style_guide)
                    return style_guide, css, logo_path
            except Exception as e:
                logger.warning(f"Failed to analyze logo: {e}, using default style")

        # Use default style
        style_guide = self.brand_analyzer._get_default_style_guide()
        css = self.brand_analyzer.generate_css_from_style(style_guide)
        return style_guide, css, None
    
    def _get_responses_queryset(self):
        """Get filtered ESRSUserResponse queryset"""
        from accounts.models import ESRSUserResponse
        from django.db.models import Q
        
        queryset = ESRSUserResponse.objects.filter(user=self.user).select_related(
            'disclosure__standard__category'
        ).order_by('disclosure__standard__order', 'disclosure__order')
        
        # Filter by standard if provided
        if self.standard_id:
            queryset = queryset.filter(disclosure__standard_id=self.standard_id)
        
        # Filter by specific disclosures if provided
        if self.disclosure_ids:
            queryset = queryset.filter(disclosure_id__in=self.disclosure_ids)
        
        # Only include responses with content
        queryset = queryset.filter(
            Q(ai_answer__isnull=False) | 
            Q(manual_answer__isnull=False) | 
            Q(final_answer__isnull=False)
        )
        
        return queryset
    
    def generate_pdf(self) -> BytesIO:
        """
        Generate branded PDF report using HTML/CSS with AI-analyzed styling
        Returns BytesIO buffer
        """
        # Get brand styling
        style_guide, custom_css, logo_path = self._get_brand_style()

        # Get responses data
        responses = self._get_responses_queryset()

        # Prepare template data
        template_data = self._prepare_template_data(responses, logo_path)

        # Render HTML from template
        template = self.jinja_env.get_template('base_report.html')
        html_content = template.render(
            custom_css=custom_css,
            **template_data
        )

        # Convert HTML to PDF using WeasyPrint
        buffer = BytesIO()
        HTML(string=html_content, base_url=str(Path(__file__).parent)).write_pdf(buffer)
        buffer.seek(0)

        logger.info(f"Generated branded PDF for user {self.user.email}")
        return buffer

    def _prepare_template_data(self, responses, logo_path: Optional[str]) -> dict:
        """Prepare data for Jinja2 template"""
        # Group responses by standard
        standards_data = []
        current_standard = None
        current_disclosures = []

        for response in responses:
            disclosure = response.disclosure
            standard = disclosure.standard

            # New standard - save previous
            if current_standard and current_standard.id != standard.id:
                standards_data.append({
                    'code': current_standard.code,
                    'name': current_standard.name,
                    'description': current_standard.description,
                    'disclosures': current_disclosures
                })
                current_disclosures = []

            current_standard = standard

            # Prepare disclosure data
            answer_text = response.final_answer or response.ai_answer or response.manual_answer
            if answer_text:
                # Convert plain text to HTML paragraphs
                answer_html = ''.join([f'<p>{para.strip()}</p>' for para in answer_text.split('\n\n') if para.strip()])
            else:
                answer_html = '<p><em>No response available</em></p>'

            disclosure_data = {
                'code': disclosure.code,
                'name': disclosure.name,
                'requirement_text': disclosure.requirement_text[:500] + '...' if len(disclosure.requirement_text) > 500 else disclosure.requirement_text,
                'answer': answer_html,
                'charts': response.chart_data[:3] if response.chart_data else [],
                'tables': response.table_data[:2] if response.table_data else []
            }

            current_disclosures.append(disclosure_data)

        # Add last standard
        if current_standard:
            standards_data.append({
                'code': current_standard.code,
                'name': current_standard.name,
                'description': current_standard.description,
                'disclosures': current_disclosures
            })

        # If no responses, create empty state
        if not standards_data:
            standards_data = [{
                'code': 'N/A',
                'name': 'No Data Available',
                'description': 'No ESRS responses are available for this report.',
                'disclosures': []
            }]

        return {
            'report_title': 'ESRS Sustainability Report',
            'report_subtitle': 'European Sustainability Reporting Standards',
            'report_type': 'European Sustainability Reporting Standards',
            'company_name': self.user.email,
            'generation_date': datetime.now().strftime('%B %d, %Y at %H:%M'),
            'logo_path': logo_path if logo_path and os.path.exists(logo_path) else None,
            'standards': standards_data
        }
    
    def generate_word(self) -> BytesIO:
        """
        Generate Word document report
        Returns BytesIO buffer
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title page
        title = doc.add_heading('ESRS Sustainability Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company info
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_para.add_run(f'Company: {self.user.email}\n').bold = True
        company_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}\n')
        company_para.add_run('Report Type: European Sustainability Reporting Standards')
        
        doc.add_page_break()
        
        # Get responses
        responses = self._get_responses_queryset()
        
        if not responses.exists():
            doc.add_paragraph('No ESRS responses available for this report.')
        else:
            current_standard = None
            
            for response in responses:
                disclosure = response.disclosure
                standard = disclosure.standard
                
                # Add standard header if changed
                if current_standard != standard.code:
                    current_standard = standard.code
                    doc.add_heading(f"{standard.code}: {standard.name}", level=1)
                
                # Disclosure header
                doc.add_heading(f"{disclosure.code}: {disclosure.name}", level=2)
                
                # Requirement
                req_para = doc.add_paragraph()
                req_para.add_run('Requirement: ').italic = True
                req_para.add_run(disclosure.requirement_text[:300] + '...')
                
                # Answer
                answer_text = response.final_answer or response.ai_answer or response.manual_answer
                if answer_text:
                    doc.add_heading('Answer:', level=3)
                    paragraphs = answer_text.split('\n\n')
                    for para in paragraphs[:5]:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                # Add charts
                if response.chart_data:
                    doc.add_heading('Visual Analytics:', level=3)
                    
                    for chart in response.chart_data[:3]:
                        try:
                            # Decode base64 image
                            image_data = base64.b64decode(chart['image_base64'])
                            img_buffer = BytesIO(image_data)
                            
                            # Add image to Word
                            doc.add_picture(img_buffer, width=Inches(5))
                            
                            # Caption
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption.add_run(f"{chart['title']} ({chart['type'].upper()} chart)")
                            caption_run.italic = True
                        except Exception as e:
                            logger.warning(f"Failed to add chart to Word: {e}")
                
                # Add tables
                if response.table_data:
                    for table_info in response.table_data[:2]:
                        try:
                            doc.add_paragraph(table_info['title'], style='Heading 3')
                            
                            # Create table
                            table = doc.add_table(
                                rows=1 + len(table_info['rows']),
                                cols=len(table_info['headers'])
                            )
                            table.style = 'Light Grid Accent 1'
                            
                            # Header row
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(table_info['headers']):
                                header_cells[i].text = str(header)
                                header_cells[i].paragraphs[0].runs[0].bold = True
                            
                            # Data rows
                            for row_idx, row_data in enumerate(table_info['rows']):
                                row_cells = table.rows[row_idx + 1].cells
                                for col_idx, cell_value in enumerate(row_data):
                                    row_cells[col_idx].text = str(cell_value)
                        except Exception as e:
                            logger.warning(f"Failed to add table to Word: {e}")
                
                doc.add_page_break()
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
